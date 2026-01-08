#!/usr/bin/env python3
"""
Word to Jekyll Post Converter
Converts .docx files to Jekyll-formatted Markdown blog posts

Usage:
    python word_to_post.py your-document.docx
    
Requirements:
    pip install python-docx
"""

import sys
import os
import re
from datetime import datetime
from pathlib import Path

try:
    from docx import Document
except ImportError:
    print("Error: python-docx is not installed.")
    print("Install it with: pip install python-docx")
    sys.exit(1)


def sanitize_filename(text):
    """Convert text to a valid filename slug"""
    text = text.lower()
    text = re.sub(r'[^a-z0-9\s-]', '', text)
    text = re.sub(r'\s+', '-', text)
    text = re.sub(r'-+', '-', text)
    return text.strip('-')


def extract_images_from_docx(doc, output_dir):
    """Extract images from Word document"""
    images = []
    image_dir = Path(output_dir)
    image_dir.mkdir(parents=True, exist_ok=True)
    
    # Get images from document relationships
    for rel in doc.part.rels.values():
        if "image" in rel.target_ref:
            try:
                image_data = rel.target_part.blob
                image_ext = rel.target_ref.split('.')[-1]
                image_name = f"image-{len(images) + 1}.{image_ext}"
                image_path = image_dir / image_name
                
                with open(image_path, 'wb') as f:
                    f.write(image_data)
                
                images.append(image_name)
                print(f"  ✓ Extracted image: {image_name}")
            except Exception as e:
                print(f"  ✗ Failed to extract image: {e}")
    
    return images


def convert_paragraph_to_markdown(paragraph):
    """Convert a Word paragraph to Markdown format"""
    text = paragraph.text.strip()
    
    if not text:
        return ""
    
    # Check paragraph style
    style = paragraph.style.name.lower()
    
    # Headings
    if 'heading 1' in style or 'title' in style:
        return f"# {text}\n"
    elif 'heading 2' in style:
        return f"## {text}\n"
    elif 'heading 3' in style:
        return f"### {text}\n"
    elif 'heading 4' in style:
        return f"#### {text}\n"
    
    # Process inline formatting
    markdown_text = ""
    for run in paragraph.runs:
        run_text = run.text
        
        # Bold
        if run.bold:
            run_text = f"**{run_text}**"
        
        # Italic
        if run.italic:
            run_text = f"*{run_text}*"
        
        # Code (if monospace font)
        if run.font.name and 'courier' in run.font.name.lower():
            run_text = f"`{run_text}`"
        
        markdown_text += run_text
    
    # Lists
    if text.startswith('•') or text.startswith('-') or text.startswith('*'):
        markdown_text = f"- {markdown_text.lstrip('•-* ')}"
    
    return f"{markdown_text}\n"


def convert_docx_to_markdown(docx_path):
    """Convert Word document to Markdown content"""
    doc = Document(docx_path)
    markdown_content = []
    
    for paragraph in doc.paragraphs:
        md_para = convert_paragraph_to_markdown(paragraph)
        if md_para:
            markdown_content.append(md_para)
    
    # Add spacing between paragraphs
    content = "\n".join(markdown_content)
    
    # Clean up multiple newlines
    content = re.sub(r'\n{3,}', '\n\n', content)
    
    return content


def create_jekyll_post(docx_path, title=None, tags=None, excerpt=None, date=None):
    """Create a Jekyll blog post from a Word document"""
    
    print(f"\n{'='*60}")
    print(f"Converting: {docx_path}")
    print(f"{'='*60}\n")
    
    # Load document
    doc = Document(docx_path)
    
    # Extract title from first paragraph if not provided
    if not title and doc.paragraphs:
        title = doc.paragraphs[0].text.strip()
        if len(title) > 100:
            title = title[:100] + "..."
    
    if not title:
        title = Path(docx_path).stem
    
    # Get date
    if not date:
        date = datetime.now().strftime('%Y-%m-%d')
    
    # Generate filename and permalink
    slug = sanitize_filename(title)
    filename = f"{date}-{slug}.md"
    year, month = date.split('-')[0], date.split('-')[1]
    permalink = f"/posts/{year}/{month}/{slug}/"
    
    # Extract images
    image_output_dir = Path('images/posts')
    images = extract_images_from_docx(doc, image_output_dir)
    
    # Convert content to markdown
    print("Converting content to Markdown...")
    content = convert_docx_to_markdown(docx_path)
    
    # Create frontmatter
    frontmatter = f"""---
title: '{title}'
date: {date}
permalink: {permalink}"""
    
    if tags:
        frontmatter += "\ntags:"
        for tag in tags:
            frontmatter += f"\n  - {tag.strip()}"
    
    if excerpt:
        frontmatter += f"\nexcerpt: '{excerpt}'"
    
    frontmatter += "\n---\n\n"
    
    # Combine frontmatter and content
    full_post = frontmatter + content
    
    # Save to _posts directory
    output_path = Path('_posts') / filename
    output_path.parent.mkdir(parents=True, exist_ok=True)
    
    with open(output_path, 'w', encoding='utf-8') as f:
        f.write(full_post)
    
    # Summary
    print(f"\n{'='*60}")
    print(f"✓ Post created successfully!")
    print(f"{'='*60}")
    print(f"File: {output_path}")
    print(f"Title: {title}")
    print(f"Date: {date}")
    print(f"Permalink: {permalink}")
    
    if images:
        print(f"\nImages extracted ({len(images)}):")
        for img in images:
            print(f"  - /images/posts/{img}")
        print(f"\nNote: Images are saved in {image_output_dir}/")
        print("You may need to manually insert them in the post.")
    
    print(f"\nNext steps:")
    print(f"1. Review and edit: {output_path}")
    print(f"2. Preview: http://localhost:4000{permalink}")
    print(f"3. Commit: git add {output_path}")
    if images:
        print(f"4. Commit images: git add images/posts/")
    print(f"5. Push: git commit -m 'Add: {title}' && git push")
    print(f"\n{'='*60}\n")
    
    return output_path


def main():
    """Main function"""
    if len(sys.argv) < 2:
        print("Usage: python word_to_post.py your-document.docx")
        print("\nOptional arguments:")
        print("  --title 'Post Title'")
        print("  --tags 'tag1,tag2,tag3'")
        print("  --excerpt 'Brief description'")
        print("  --date 'YYYY-MM-DD'")
        print("\nExample:")
        print("  python word_to_post.py research.docx --title 'My Research' --tags 'AI,ML'")
        sys.exit(1)
    
    docx_path = sys.argv[1]
    
    if not os.path.exists(docx_path):
        print(f"Error: File '{docx_path}' not found!")
        sys.exit(1)
    
    if not docx_path.endswith('.docx'):
        print("Error: File must be a .docx Word document")
        sys.exit(1)
    
    # Parse optional arguments
    title = None
    tags = []
    excerpt = None
    date = None
    
    i = 2
    while i < len(sys.argv):
        if sys.argv[i] == '--title' and i + 1 < len(sys.argv):
            title = sys.argv[i + 1]
            i += 2
        elif sys.argv[i] == '--tags' and i + 1 < len(sys.argv):
            tags = [t.strip() for t in sys.argv[i + 1].split(',')]
            i += 2
        elif sys.argv[i] == '--excerpt' and i + 1 < len(sys.argv):
            excerpt = sys.argv[i + 1]
            i += 2
        elif sys.argv[i] == '--date' and i + 1 < len(sys.argv):
            date = sys.argv[i + 1]
            i += 2
        else:
            i += 1
    
    # Convert the document
    try:
        create_jekyll_post(docx_path, title, tags, excerpt, date)
    except Exception as e:
        print(f"Error: {e}")
        import traceback
        traceback.print_exc()
        sys.exit(1)


if __name__ == "__main__":
    main()
